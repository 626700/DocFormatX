"""
文档渲染器 - 将样式 + 内容块合成为 Word 文档。
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import logging

from .style_parser import StyleDocument, TextStyle
from .content_parser import ContentBlock, parse_content

log = logging.getLogger(__name__)

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _apply_text_style(paragraph, style: TextStyle):
    for run in paragraph.runs:
        if style.font_name:
            run.font.name = style.font_name
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:eastAsia'), style.font_name)
        if style.size_pt:
            run.font.size = Pt(style.size_pt)
        if style.bold is not None:
            run.bold = style.bold


def _apply_paragraph_style(paragraph, style: TextStyle):
    pf = paragraph.paragraph_format
    if style.alignment and style.alignment in ALIGN_MAP:
        paragraph.alignment = ALIGN_MAP[style.alignment]
    if style.space_before_pt is not None:
        pf.space_before = Pt(style.space_before_pt)
    if style.space_after_pt is not None:
        pf.space_after = Pt(style.space_after_pt)
    if style.line_spacing is not None:
        pf.line_spacing = style.line_spacing
    if style.first_line_indent_chars is not None:
        indent_pt = (style.size_pt or 12) * style.first_line_indent_chars
        pf.first_line_indent = Pt(indent_pt)


def render(style_path: Path, content_path: Path, output_path: Path):
    style_doc = StyleDocument.from_yaml(style_path)
    log.info("已加载样式: %s", style_doc.name)

    blocks = parse_content(content_path)
    log.info("已解析 %d 个内容块", len(blocks))

    doc = Document()
    default_font = style_doc.body.font_name or "宋体"
    style = doc.styles["Normal"]
    style.font.name = default_font
    style.element.rPr.rFonts.set(qn('w:eastAsia'), default_font)

    # 目录
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.append(fldChar_begin)
    instrText = run._element.makeelement(qn('w:instrText'), {})
    instrText.text = " TOC \\o \"1-3\" \\h \\z \\u "
    run._element.append(instrText)
    fldChar_separate = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run._element.append(fldChar_separate)
    run2 = paragraph.add_run("（请在 Word 中右键此处 → 更新域 以生成目录）")
    run2.font.size = Pt(10)
    fldChar_end = run2._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run2._element.append(fldChar_end)
    doc.add_paragraph()

    # 逐块写入
    for block in blocks:
        if block.type == "title":
            p = doc.add_paragraph(block.text, style="Title")
            _apply_paragraph_style(p, style_doc.title)
            _apply_text_style(p, style_doc.title)

        elif block.type == "heading_1":
            p = doc.add_paragraph(block.text, style="Heading 1")
            _apply_paragraph_style(p, style_doc.heading_1)
            _apply_text_style(p, style_doc.heading_1)

        elif block.type == "heading_2":
            p = doc.add_paragraph(block.text, style="Heading 2")
            _apply_paragraph_style(p, style_doc.heading_2)
            _apply_text_style(p, style_doc.heading_2)

        elif block.type == "body":
            p = doc.add_paragraph(block.text)
            _apply_paragraph_style(p, style_doc.body)
            _apply_text_style(p, style_doc.body)

        elif block.type == "table":
            lines = block.text.strip().split("\n")
            rows = []
            for line in lines:
                line = line.strip()
                if line.startswith("|") and line.endswith("|"):
                    cells = [c.strip() for c in line[1:-1].split("|")]
                    rows.append(cells)
            if not rows:
                continue

            transpose = block.props.get("方向", "") == "竖"
            border_style_key = block.props.get("样式", style_doc.table_style.default_style)
            three_line = border_style_key in ("三线表", "three_line")

            if transpose:
                header = rows[0]
                data = rows[1:]
                data_T = list(map(list, zip(*data)))
                rows = [[header[i]] + data_T[i] for i in range(len(header))]

            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Light Grid Accent 1'

            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    cell = table.cell(i, j)
                    cell.text = cell_text
                    if i == 0 and style_doc.table_style.header_bold:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

            # 边框
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn('w:tblPr'))
            for old_borders in tblPr.findall(qn('w:tblBorders')):
                tblPr.remove(old_borders)
            borders = etree.SubElement(tblPr, qn('w:tblBorders'))

            if three_line:
                for side, sz in [('top', '24'), ('bottom', '24'), ('insideH', '8')]:
                    elem = etree.SubElement(borders, qn(f'w:{side}'))
                    elem.set(qn('w:val'), 'single')
                    elem.set(qn('w:sz'), sz)
                    elem.set(qn('w:space'), '0')
                    elem.set(qn('w:color'), '000000')
                for side in ['left', 'right', 'insideV']:
                    elem = etree.SubElement(borders, qn(f'w:{side}'))
                    elem.set(qn('w:val'), 'none')
                    elem.set(qn('w:sz'), '0')
                    elem.set(qn('w:space'), '0')
                    elem.set(qn('w:color'), 'auto')
            else:
                yb = style_doc.table_style.borders
                for side in ['top', 'bottom', 'insideH', 'insideV', 'left', 'right']:
                    rule = getattr(yb, side)
                    elem = etree.SubElement(borders, qn(f'w:{side}'))
                    if rule.visible:
                        elem.set(qn('w:val'), 'single')
                        elem.set(qn('w:sz'), str(int(rule.weight_pt * 8)))
                        elem.set(qn('w:space'), '0')
                        elem.set(qn('w:color'), rule.color)
                    else:
                        elem.set(qn('w:val'), 'none')
                        elem.set(qn('w:sz'), '0')
                        elem.set(qn('w:space'), '0')
                        elem.set(qn('w:color'), 'auto')
            doc.add_paragraph()

        elif block.type == "image":
            img_path = Path(block.image_path)
            if img_path.exists():
                p = doc.add_paragraph()
                if style_doc.inline_image.alignment in ALIGN_MAP:
                    p.alignment = ALIGN_MAP[style_doc.inline_image.alignment]
                run = p.add_run()
                run.add_picture(str(img_path), width=Cm(style_doc.inline_image.width_cm))
            else:
                log.warning("图片不存在: %s", img_path)
                doc.add_paragraph(f"[图片缺失: {img_path.name}]")

    # 页眉页脚
    section_rules = style_doc.sections
    if section_rules:
        while len(doc.sections) < len(section_rules):
            doc.add_section()
        for i, rule in enumerate(section_rules):
            section = doc.sections[i]
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            if rule.header_text:
                header = section.header
                hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                hp.text = rule.header_text
            if rule.page_number_start is not None:
                footer = section.footer
                fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = fp.add_run()
                fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
                run._element.append(fldChar1)
                instrText = run._element.makeelement(qn('w:instrText'), {})
                instrText.text = " PAGE "
                run._element.append(instrText)
                fldChar2 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
                run._element.append(fldChar2)
                sectPr = section._sectPr
                pgNumType = sectPr.find(qn('w:pgNumType'))
                if pgNumType is None:
                    pgNumType = etree.SubElement(sectPr, qn('w:pgNumType'))
                pgNumType.set(qn('w:start'), str(rule.page_number_start))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info("文档已生成: %s", output_path)
    return output_path